"""Step 8: DOCX extraction — heading-style sectioning, tables, core properties, fallbacks.

Synthetic .docx files built with python-docx so the tests are deterministic and committable.
"""

from datetime import datetime
from pathlib import Path

import docx as _docx
import pytest

from locus.extract.base import MAX_SECTION_CHARS
from locus.extract.docx import extract_docx

BODY = " ".join(f"Sentence {i} contains several descriptive words about the topic." for i in range(12))


def _make_docx(path: Path, *, title=None, created=None) -> "_docx.document.Document":
    document = _docx.Document()
    if title is not None:
        document.core_properties.title = title
    if created is not None:
        document.core_properties.created = created
    document._locus_path = path  # stashed for _save
    return document


def _save(document) -> Path:
    path = document._locus_path
    document.save(str(path))
    return path


def test_heading_styles_split_sections(tmp_path: Path):
    d = _make_docx(tmp_path / "report.docx", title="Project Phoenix Write-up",
                   created=datetime(2024, 3, 9, 12, 0, 0))
    d.add_heading("Background", level=1)
    d.add_paragraph(BODY)
    d.add_heading("Implementation", level=2)
    d.add_paragraph(BODY)
    doc = extract_docx(_save(d))
    assert doc.title == "Project Phoenix Write-up"
    assert doc.source_date == "2024-03-09"
    assert doc.section_strategy == "headings"
    assert doc.page_count == 0
    assert [s.title for s in doc.sections] == ["Background", "Implementation"]
    assert doc.sections[0].text.startswith("Background")  # heading kept in verbatim text


def test_table_text_is_retrievable(tmp_path: Path):
    d = _make_docx(tmp_path / "table.docx")
    d.add_heading("Results", level=1)
    d.add_paragraph(BODY)
    table = d.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Metric"
    table.cell(0, 1).text = "Value"
    table.cell(1, 0).text = "Sharpe ratio"
    table.cell(1, 1).text = "1.42"
    doc = extract_docx(_save(d))
    joined = "\n".join(s.text for s in doc.sections)
    assert "Sharpe ratio | 1.42" in joined


def test_no_heading_styles_falls_back_to_single_or_windowed(tmp_path: Path):
    d = _make_docx(tmp_path / "flat.docx")
    for _ in range(3):
        d.add_paragraph(BODY)
    doc = extract_docx(_save(d))
    assert doc.section_strategy == "single"
    assert len(doc.sections) == 1
    assert doc.sections[0].title is None

    d2 = _make_docx(tmp_path / "flat-long.docx")
    for _ in range(40):  # well over MAX_SECTION_CHARS in total
        d2.add_paragraph(BODY)
    doc2 = extract_docx(_save(d2))
    assert doc2.section_strategy == "windowed"
    assert len(doc2.sections) > 1
    assert all(len(s.text) <= MAX_SECTION_CHARS for s in doc2.sections)


def test_title_falls_back_to_first_heading_then_stem(tmp_path: Path):
    d = _make_docx(tmp_path / "untitled.docx")  # no core-properties title
    d.add_heading("Actual Document Heading", level=1)
    d.add_paragraph(BODY)
    doc = extract_docx(_save(d))
    assert doc.title == "Actual Document Heading"

    d2 = _make_docx(tmp_path / "plain-paras.docx")
    d2.add_paragraph(BODY)
    doc2 = extract_docx(_save(d2))
    assert doc2.title == "plain-paras"


def test_empty_docx_raises(tmp_path: Path):
    d = _make_docx(tmp_path / "empty.docx")
    with pytest.raises(ValueError, match="no extractable text"):
        extract_docx(_save(d))


def test_non_docx_raises_value_error(tmp_path: Path):
    bogus = tmp_path / "bogus.docx"
    bogus.write_bytes(b"this is not a zip archive")
    with pytest.raises(ValueError, match="not a readable docx"):
        extract_docx(bogus)

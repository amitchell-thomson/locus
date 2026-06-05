"""Extract structured content from a .docx file with python-docx.

Produces an `ExtractedDoc` (see extract/base.py). Section boundaries come from
heading-styled paragraphs ("Heading 1".."Heading 9", or "Title"); a document with no
heading styles at all (common for Google-Docs/word-processor exports) degrades to a
single windowed section, like plain text.

Tables are flattened to pipe-joined rows and appended after the paragraph flow.
LIMITATION (documented, not a bug): python-docx exposes paragraphs and tables as separate
collections, so table↔paragraph interleaving is approximated — each table lands at the end
of the document body rather than at its exact in-flow position. Acceptable: retrieval
grounds on chunks regardless of intra-section ordering (§15.0); revisit by walking the XML
body only if real documents show it matters.

Section text includes its own heading line (mirrors pdf.py / textdoc.py), so merged small
sections keep absorbed headings visible.
"""

from __future__ import annotations

import re
from pathlib import Path

import docx as _docx  # python-docx; aliased so the module name never reads ambiguously

from locus.extract.base import ExtractedDoc, build_sections

# Heading-styled paragraph: built-in "Heading 1".."Heading 9" (localised templates keep the
# English style *name* in the underlying XML via python-docx) or the document "Title" style.
_HEADING_STYLE = re.compile(r"^(?:Heading [1-9]|Title)$")


def extract_docx(path: str | Path) -> ExtractedDoc:
    """Extract a structured ExtractedDoc from the .docx file at `path`."""
    path = Path(path)
    try:
        document = _docx.Document(str(path))
    except Exception as exc:  # python-docx raises several types on a damaged/non-docx file
        raise ValueError(f"not a readable docx: {exc}") from exc

    spans = _split_by_headings(document)
    table_text = _tables_as_text(document)
    if table_text:
        if spans:
            title, text = spans[-1]
            spans[-1] = (title, f"{text}\n\n{table_text}")
        else:
            spans = [(None, table_text)]

    if not any(text.strip() for _, text in spans):
        raise ValueError("no extractable text")

    sections = build_sections(spans)
    if any(title is not None for title, _ in spans):
        strategy = "headings"
    elif len(sections) > 1:
        strategy = "windowed"
    else:
        strategy = "single"

    return ExtractedDoc(
        title=_resolve_title(document, spans, path),
        page_count=0,  # .docx has no fixed pagination; "page" is vestigial here
        section_strategy=strategy,
        sections=sections,
        source_path=str(path),
        source_date=_resolve_source_date(document),
    )


def _split_by_headings(document) -> list[tuple[str | None, str]]:
    """(title, text) spans split at heading-styled paragraphs; text includes the heading."""
    spans: list[tuple[str | None, str]] = []
    cur_title: str | None = None
    cur_parts: list[str] = []
    for para in document.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        style = (para.style.name if para.style is not None else "") or ""
        if _HEADING_STYLE.match(style):
            if cur_parts:
                spans.append((cur_title, "\n\n".join(cur_parts)))
            cur_title = text
            cur_parts = [text]
        else:
            cur_parts.append(text)
    if cur_parts:
        spans.append((cur_title, "\n\n".join(cur_parts)))
    return spans


def _tables_as_text(document) -> str:
    """All tables flattened to pipe-joined rows (one table per blank-line-separated block)."""
    blocks: list[str] = []
    for table in document.tables:
        rows = []
        for row in table.rows:
            cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            if any(cells):
                rows.append(" | ".join(cells))
        if rows:
            blocks.append("\n".join(rows))
    return "\n\n".join(blocks)


def _resolve_title(document, spans: list[tuple[str | None, str]], path: Path) -> str | None:
    """Core-properties title, else the first heading, else the filename stem."""
    meta = (document.core_properties.title or "").strip()
    if len(meta) > 3:
        return meta
    for title, _ in spans:
        if title:
            return title
    return path.stem


def _resolve_source_date(document) -> str | None:
    """Core-properties created date if present, else modified date, else None."""
    props = document.core_properties
    for stamp in (props.created, props.modified):
        if stamp is not None:
            return stamp.date().isoformat()
    return None
